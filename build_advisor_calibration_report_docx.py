#!/usr/bin/env python3
"""Build an advisor-facing calibration report DOCX."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CN_FONT = "宋体"
EN_FONT = "Times New Roman"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
INK = "000000"
CHANNELS = ("r", "g", "b")
GITHUB_URL = "https://github.com/fyhloveyh-afk/radiometric-calibration-rgb"


def find_root() -> Path:
    return next(Path.cwd().rglob("CalibrationCapture_20260522-23_800-1500_12bit16"))


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size: float, bold: bool = False, color: str = INK) -> None:
    style.font.name = EN_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, width_dxa: int = 9000, indent_dxa: int = 0) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.6)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    styles = doc.styles
    set_style_font(styles["Normal"], 10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 6),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ]:
        set_style_font(styles[name], size, bold=True, color=color)
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)
        styles[name].paragraph_format.line_spacing = 1.25

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("十相机黑体炉辐射标定方法与结果报告")
    set_run_font(run, 9, color="666666")

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—  ")
    set_run_font(run, 9, color="777777")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = footer.add_run("  —")
    set_run_font(run2, 9, color="777777")


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("十相机黑体炉辐射标定方法与结果报告")
    set_run_font(r, 22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAW 与 BMP 两类数据的标定流程、质量控制与结果对比")
    set_run_font(r, 14, color="333333")
    p.paragraph_format.space_after = Pt(48)

    meta = [
        ("数据范围", "800-1500°C 黑体炉标定数据"),
        ("相机数量", "10 个相机，每个工况理论重复 3 次"),
        ("主要模型", "standard poly3：L_eff = aX^3 + bX^2 + cX + d"),
        ("代码地址", GITHUB_URL),
        ("报告日期", date.today().isoformat()),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, width_dxa=7200)
    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_width(table.cell(i, 0), 1800)
        set_cell_width(table.cell(i, 1), 5400)
        set_cell_shading(table.cell(i, 0), LIGHT_GRAY)
        for cell in table.rows[i].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell == table.cell(i, 0) else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, 10.5, bold=(cell == table.cell(i, 0)))
    doc.add_page_break()


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, 10.5)


def add_noindent(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 10.5)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, 10.5)


def fmt(value) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (float, int)):
        value = float(value)
        if value == 0:
            return "0"
        if abs(value) < 0.0001 or abs(value) >= 10000:
            return f"{value:.3e}"
        return f"{value:.6f}".rstrip("0").rstrip(".")
    return str(value)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int] | None = None, font_size: float = 9) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], LIGHT_GRAY)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = fmt(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[i].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, font_size)
    if widths:
        for table_row in table.rows:
            for i, width in enumerate(widths):
                set_cell_width(table_row.cells[i], width)
    doc.add_paragraph()


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], widths: list[int] | None = None, font_size: float = 9) -> None:
    add_table(doc, headers, df[columns].values.tolist(), widths, font_size)


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 9)
    r.italic = True
    cap.paragraph_format.space_after = Pt(8)


def read_coefficients(run_dir: Path, corrected: bool = False) -> pd.DataFrame:
    filename = "calibration_coefficients_intercept_corrected.json" if corrected else "calibration_coefficients_standard.json"
    rows = []
    for cam in range(1, 11):
        cam_name = f"camera{cam:02d}"
        data = json.loads((run_dir / cam_name / filename).read_text(encoding="utf-8"))
        for ch in CHANNELS:
            coefs = data[ch]["coefficients"]
            metrics = data[ch]["metrics"]
            rows.append(
                {
                    "相机": cam_name,
                    "通道": ch.upper(),
                    "a": coefs[0],
                    "b": coefs[1],
                    "c": coefs[2],
                    "d": coefs[3],
                    "RMSE": metrics["rmse"],
                    "R2": metrics["r2"],
                }
            )
    return pd.DataFrame(rows)


def main() -> None:
    root = find_root()
    asset_dir = root / "calibration_report_assets"
    raw_final = root / "calibration_fit_10cam_auto_roi_badpixel_poly3_saturation_excluded"
    bmp_final = root / "calibration_fit_10cam_bmp_auto_roi_poly3_saturation_excluded"
    sat_dir = root / "saturation_diagnostics_auto_roi"
    out_dir = root / "calibration_report"
    out_dir.mkdir(exist_ok=True)
    out_docx = out_dir / "十相机黑体炉辐射标定方法与结果报告_导师版.docx"

    raw_summary = pd.read_csv(raw_final / "saturation_exclusion_metric_summary.csv")
    bmp_summary = pd.read_csv(bmp_final / "saturation_exclusion_metric_summary.csv")
    raw_counts = pd.read_csv(raw_final / "saturation_filter_row_counts.csv")
    bmp_counts = pd.read_csv(bmp_final / "saturation_filter_row_counts.csv")
    bias_summary = pd.read_csv(sat_dir / "standard_vs_intercept_corrected_summary.csv")
    raw_bias = bias_summary[bias_summary["run"] == "RAW saturation excluded"].copy()
    bmp_bias = bias_summary[bias_summary["run"] == "BMP saturation excluded"].copy()

    doc = Document()
    setup_document(doc)
    add_title_page(doc)

    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "本报告总结了十相机黑体炉辐射标定的完整处理方法与结果。标定数据覆盖 800-1500°C，"
        "每个工况理论采集 3 次重复图像。处理流程包括两天数据合并、计划工况筛选、重复采集容错、"
        "自动 ROI 定位、暗场校正、坏点屏蔽、过曝工况剔除以及 RAW/BMP 两类数据的拟合对比。"
    )
    add_para(
        doc,
        "最终推荐以 RAW 数据作为主标定结果，采用 standard poly3 模型。过曝工况剔除后，RAW 分支 R/G 通道拟合质量显著提升，"
        "R 通道平均 RMSE 由 0.065477 降至 0.024874，G 通道平均 RMSE 由 0.015731 降至 0.010686；"
        "B 通道 R2 仍保持在 0.9987 以上。BMP 分支作为显示值标定方案也可用，其结果与 RAW 主线趋势一致。"
    )

    doc.add_heading("一、实验数据与标定目标", level=1)
    add_para(
        doc,
        "标定对象为 10 个相机的 R/G/B 三通道响应。实验中黑体炉温度范围为 800-1500°C，"
        "不同温度下设置了多组曝光时间。由于数据分两天采集，原始数据由软件分为多个文件夹，"
        "后续处理首先将计划内标定工况合并到统一数据结构中。"
    )
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["温度范围", "800-1500°C"],
            ["相机数量", "10 个相机"],
            ["通道", "R、G、B 三通道"],
            ["重复采集", "每个工况理论 repeat1、repeat2、repeat3 三次"],
            ["明场数据", "两天采集的黑体炉图像，合并后按计划工况筛选"],
            ["暗场数据", "Camera01 暗场，覆盖全部设计曝光时间；当前作为十相机统一暗场参考"],
        ],
        widths=[2200, 6800],
    )
    add_para(
        doc,
        "标定目标是建立每个相机、每个颜色通道的辐射响应函数，将曝光归一化后的相机响应量 X 映射为该通道的等效黑体辐亮度 L_eff。"
    )

    doc.add_heading("二、辐射参考值的计算", level=1)
    add_para(
        doc,
        "纵坐标 L_eff 不是相机灰度值，也不是温度，而是由黑体温度通过 Planck 公式计算得到的通道等效辐亮度。"
        "脚本首先计算黑体在各波长处的光谱辐亮度，然后乘以相机通道光谱响应和 650 nm short-pass 滤光片透过率，"
        "最后在波长维度上积分。"
    )
    add_equation(doc, "S_eff,c(λ) = S_camera,c(λ) · T_filter(λ)")
    add_equation(doc, "L_eff,c = ∫ L_λ(T) S_eff,c(λ)dλ / ∫ S_eff,c(λ)dλ")
    add_para(
        doc,
        "本次使用 normalized 参考方式，因此 L_eff 的单位为 W/(m²·sr·nm)，可以理解为经过通道响应加权后的等效光谱辐亮度。"
    )

    doc.add_heading("三、数据预处理方法", level=1)
    doc.add_heading("3.1 工况筛选与重复采集容错", level=2)
    add_para(
        doc,
        "首先根据标定工况表筛选计划内温度/曝光组合，并删除额外采集的 repeat4。"
        "对于少量缺失 repeat 的工况，不直接删除整个条件，而是在至少保留两次重复采集的前提下用可用 repeat 求平均。"
        "这样可以减少偶发漏拍对标定覆盖范围的影响。"
    )

    doc.add_heading("3.2 自动 ROI 定位", level=2)
    add_para(
        doc,
        "由于切换相机时画面位置存在偏移，固定 ROI 不能保证所有图像都落在黑体炉内部。"
        "本次采用 auto-anchor 策略：先在高温图像中检测稳定热点中心，建立锚点；随后每张图在锚点附近局部搜索黑体炉区域。"
        "对于低温高曝光导致环境也变亮的情况，若局部检测置信度不足，则回退到锚点 ROI。"
    )

    doc.add_heading("3.3 暗场与坏点处理", level=2)
    add_para(
        doc,
        "RAW 分支使用同曝光暗场均值进行暗场校正。由于暗场只采集了 Camera01，本次将 Camera01 暗场作为十相机统一暗场。"
        "同时根据暗场 RAW median 图构建坏点 mask，阈值设为 8192 DN，ROI 求均值时忽略这些异常像素。"
        "结果表明，在 120×120 ROI 下坏点处理对整体拟合指标影响很小，但能够提高长曝光暗场下的鲁棒性。"
    )
    add_para(
        doc,
        "BMP 分支默认不减暗场。前期检查表明 BMP 暗场主体灰度约为 2-3 DN，相比有效信号较小，直接拟合更简洁；"
        "该分支应理解为 BMP 显示值标定，而不是 RAW 物理响应标定。"
    )

    doc.add_heading("四、过曝诊断与剔除", level=1)
    add_para(
        doc,
        "复查过程中发现，各温度下高曝光尾部存在明显饱和。过曝会使响应曲线高端被截断，导致拟合曲线被非物理数据拉偏。"
        "因此，本次最终拟合前对实际 auto ROI 内的每个相机、每个通道、每个 repeat 进行饱和检查。"
    )
    add_table(
        doc,
        ["类别", "判据", "处理方式"],
        [
            ["明确过曝", "ROI 内任一通道/重复满足 sat_pct > 0.01% 或 near_pct > 0.5%", "从拟合中剔除"],
            ["接近饱和", "未达到剔除阈值，但 near_pct > 0.05% 或高分位数接近上限", "保留并标记复查"],
        ],
        widths=[1800, 5100, 2100],
    )
    add_table(
        doc,
        ["数据分支", "平均原始工况数", "平均保留工况数", "平均剔除工况数"],
        [
            ["RAW", raw_counts["input_rows"].mean(), raw_counts["kept_rows"].mean(), raw_counts["excluded_rows"].mean()],
            ["BMP", bmp_counts["input_rows"].mean(), bmp_counts["kept_rows"].mean(), bmp_counts["excluded_rows"].mean()],
        ],
        widths=[2200, 2300, 2300, 2200],
    )
    add_picture(doc, asset_dir / "saturation_filter_counts.png", "图 1  十相机过曝剔除前后的工况数量统计")
    add_picture(doc, asset_dir / "saturation_exclusion_metric_comparison.png", "图 2  过曝剔除前后 RAW/BMP 拟合指标变化")

    doc.add_heading("五、拟合模型与模型选择", level=1)
    add_para(
        doc,
        "比较 linear、poly2、poly3、logpoly2 后，三次多项式在当前数据上取得了最好的综合指标。"
        "最终主模型采用 standard poly3："
    )
    add_equation(doc, "X_c = (DN_c - Dark_c) / exposure_s")
    add_equation(doc, "L_eff,c = a_cX_c³ + b_cX_c² + c_cX_c + d_c")
    add_para(
        doc,
        "BMP 分支对应的输入为 X_c = DN_bmp,c / exposure_s。"
        "残差偏置修正也进行了对比：RAW 分支中 R/G 通道 standard 更好，B 通道修正略好但幅度很小；"
        "BMP 分支修正后指标略有改善。考虑物理含义和主标定稳定性，RAW 主结果推荐不做残差偏置修正。"
    )
    add_df_table(
        doc,
        raw_summary,
        ["channel", "old_rmse_mean", "new_rmse_mean", "old_r2_mean", "new_r2_mean"],
        ["通道", "剔除前 RMSE", "剔除后 RMSE", "剔除前 R²", "剔除后 R²"],
        widths=[1200, 2100, 2100, 1800, 1800],
    )
    add_df_table(
        doc,
        bmp_summary,
        ["channel", "old_rmse_mean", "new_rmse_mean", "old_r2_mean", "new_r2_mean"],
        ["通道", "BMP 剔除前 RMSE", "BMP 剔除后 RMSE", "BMP 剔除前 R²", "BMP 剔除后 R²"],
        widths=[1200, 2200, 2200, 1700, 1700],
    )

    doc.add_heading("六、十相机标定结果", level=1)
    doc.add_heading("6.1 RAW 主标定结果", level=2)
    add_para(
        doc,
        "RAW 主标定在剔除过曝条件后，R/G 通道拟合质量显著提高。"
        "从拟合曲线看，十个相机在三个通道上均表现出一致的单调响应趋势，说明过曝剔除后数据更符合辐射响应规律。"
    )
    add_picture(doc, asset_dir / "raw_saturation_excluded_fit_overview_part1.png", "图 3  RAW 主标定拟合结果：Camera01-Camera05")
    add_picture(doc, asset_dir / "raw_saturation_excluded_fit_overview_part2.png", "图 4  RAW 主标定拟合结果：Camera06-Camera10")

    doc.add_heading("6.2 BMP 分支结果", level=2)
    add_para(
        doc,
        "BMP 分支使用 8-bit BMP 图像的 ROI 均值直接拟合，不减暗场。"
        "其结果与 RAW 分支趋势一致，R/G 通道在过曝剔除后同样明显改善，说明 BMP 分支可作为显示链路或预览图像的经验标定方案。"
    )
    add_picture(doc, asset_dir / "bmp_saturation_excluded_fit_overview_part1.png", "图 5  BMP 分支拟合结果：Camera01-Camera05")
    add_picture(doc, asset_dir / "bmp_saturation_excluded_fit_overview_part2.png", "图 6  BMP 分支拟合结果：Camera06-Camera10")

    doc.add_page_break()
    doc.add_heading("七、脚本实现与复现方式", level=1)
    add_para(
        doc,
        "本次标定脚本按“数据整理—测量值提取—质量控制—模型拟合—结果导出”的顺序实现。"
        "为了保证两天采集数据、缺失 repeat、过曝工况和相机视场偏移都能被统一处理，脚本没有直接依赖固定文件数量，"
        "而是以标定工况表、manifest 信息和实际图像文件共同确定有效样本。"
    )
    add_table(
        doc,
        ["脚本", "主要功能"],
        [
            [
                "images_to_blackbody_measurements.py",
                "读取标定工况表和采集 manifest，合并 RAW/BMP 图像；完成 repeat 容错、自动 ROI、暗场校正、坏点屏蔽和过曝筛选；输出每个相机、温度、曝光、通道的平均响应值。",
            ],
            [
                "radiometric_calibration_rgb.py",
                "根据黑体温度计算通道等效辐亮度 L_eff；比较 linear、poly2、poly3、logpoly2 等模型；输出拟合系数、RMSE、R²、残差偏置修正结果和拟合图。",
            ],
            [
                "diagnose_raw_saturation.py / check_raw_overexposure.py",
                "用于独立检查 RAW 图像的饱和和接近饱和情况，辅助确定过曝剔除阈值，并生成过曝诊断统计。",
            ],
            [
                "export_monotonic_calibration.py",
                "预留用于导出单调 LUT/PCHIP 标定表，便于后续部署端采用查表或插值方式进行快速标定。",
            ],
            [
                "build_advisor_calibration_report_docx.py",
                "自动汇总标定结果、图表和系数表，生成本报告，保证后续更新数据或模型后可以重复生成同一格式的汇报文档。",
            ],
        ],
        widths=[3000, 6000],
        font_size=8.2,
    )
    add_para(
        doc,
        "整体处理流程为：首先将计划工况与实际采集文件匹配，删除 repeat4 等额外采集；随后对每张图像定位 ROI 并提取曝光归一化响应；"
        "再按相机、温度、曝光和通道聚合可用 repeat，剔除过曝样本；最后将响应量 X 与 L_eff 拟合，生成每个相机每个通道的系数文件和结果图。"
    )
    add_para(
        doc,
        "代码已整理在 GitHub 仓库中，便于导师或后续使用者查看脚本实现和复现实验流程。代码地址："
        f"{GITHUB_URL}。"
    )

    doc.add_heading("八、结论", level=1)
    add_para(
        doc,
        "本次标定的关键结论如下：第一，过曝工况必须从拟合中剔除，否则高曝光尾部会显著拉偏 R/G 通道曲线；"
        "第二，自动 ROI 能够解决相机切换造成的视场偏移，是多相机统一处理的必要步骤；"
        "第三，RAW 分支的 standard poly3 是当前最适合作为主标定的系数模型；"
        "第四，BMP 分支可作为独立的显示值标定分支，但不应与 RAW 物理响应标定混用。"
    )
    add_para(
        doc,
        "后续建议采集每个相机各自的暗场，以替代当前 Camera01 暗场复用方案；"
        "同时可在最终过曝筛选规则固化后继续导出单调 LUT/PCHIP 形式，用于部署端查表或插值。"
    )

    doc.add_heading("附录 A：RAW 主标定系数", level=1)
    raw_coef = read_coefficients(raw_final, corrected=False)
    for ch in ["R", "G", "B"]:
        doc.add_heading(f"RAW {ch} 通道", level=2)
        df = raw_coef[raw_coef["通道"] == ch]
        add_df_table(
            doc,
            df,
            ["相机", "通道", "a", "b", "c", "d", "RMSE", "R2"],
            ["相机", "通道", "a", "b", "c", "d", "RMSE", "R²"],
            widths=[1200, 900, 1250, 1250, 1250, 1250, 950, 950],
            font_size=7.5,
        )

    doc.add_page_break()
    doc.add_heading("附录 B：BMP 分支修正模型系数", level=1)
    bmp_coef = read_coefficients(bmp_final, corrected=True)
    for ch in ["R", "G", "B"]:
        doc.add_heading(f"BMP {ch} 通道", level=2)
        df = bmp_coef[bmp_coef["通道"] == ch]
        add_df_table(
            doc,
            df,
            ["相机", "通道", "a", "b", "c", "d", "RMSE", "R2"],
            ["相机", "通道", "a", "b", "c", "d", "RMSE", "R²"],
            widths=[1200, 900, 1250, 1250, 1250, 1250, 950, 950],
            font_size=7.5,
        )

    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
