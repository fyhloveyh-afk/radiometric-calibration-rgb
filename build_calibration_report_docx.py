#!/usr/bin/env python3
"""Build the 10-camera calibration technical report DOCX."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
INK = "222222"
CHANNELS = ("r", "g", "b")


def find_root() -> Path:
    return next(Path.cwd().rglob("CalibrationCapture_20260522-23_800-1500_12bit16"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(34, 34, 34)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "10 相机黑体炉标定技术报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.text = "D:\\标定 · calibration report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(120, 120, 120)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run("\n" + body)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    for line in text.splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(60, 60, 60)


def fmt_value(value, digits: int = 6) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (float, int)):
        value = float(value)
        if value == 0:
            return "0"
        if abs(value) < 0.0001 or abs(value) >= 100000:
            return f"{value:.3e}"
        return f"{value:.{digits}f}".rstrip("0").rstrip(".")
    return str(value)


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int] | None = None, font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
        set_cell_shading(hdr[i], LIGHT_GRAY)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = fmt_value(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if i > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                set_cell_width(row.cells[i], width)
    doc.add_paragraph()


def add_dataframe_table(doc: Document, df: pd.DataFrame, headers: list[str] | None = None, columns: list[str] | None = None, widths: list[int] | None = None, font_size: int = 8) -> None:
    if columns is not None:
        df = df[columns]
    display_headers = headers or list(df.columns)
    rows = df.values.tolist()
    add_table_from_rows(doc, display_headers, rows, widths=widths, font_size=font_size)


def add_picture(doc: Document, path: Path, caption: str, width_in: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def coefficient_rows(run_dir: Path, corrected: bool = False) -> pd.DataFrame:
    filename = "calibration_coefficients_intercept_corrected.json" if corrected else "calibration_coefficients_standard.json"
    rows: list[dict[str, object]] = []
    for cam in range(1, 11):
        cam_name = f"camera{cam:02d}"
        data = load_json(run_dir / cam_name / filename)
        for ch in CHANNELS:
            rec = data[ch]
            coefs = rec["coefficients"]
            metrics = rec["metrics"]
            rows.append(
                {
                    "camera": cam_name,
                    "channel": ch.upper(),
                    "a_x3": coefs[0],
                    "b_x2": coefs[1],
                    "c_x": coefs[2],
                    "d": coefs[3],
                    "rmse": metrics["rmse"],
                    "r2": metrics["r2"],
                }
            )
    return pd.DataFrame(rows)


def rel(path: Path, root: Path) -> str:
    try:
        return str(path.relative_to(root))
    except ValueError:
        return str(path)


def main() -> None:
    root = find_root()
    asset_dir = root / "calibration_report_assets"
    raw_final = root / "calibration_fit_10cam_auto_roi_badpixel_poly3_saturation_excluded"
    bmp_final = root / "calibration_fit_10cam_bmp_auto_roi_poly3_saturation_excluded"
    sat_dir = root / "saturation_diagnostics_auto_roi"
    output_dir = root / "calibration_report"
    output_dir.mkdir(exist_ok=True)
    output_docx = output_dir / "10相机黑体炉标定技术报告.docx"

    raw_summary = pd.read_csv(raw_final / "saturation_exclusion_metric_summary.csv")
    bmp_summary = pd.read_csv(bmp_final / "saturation_exclusion_metric_summary.csv")
    raw_counts = pd.read_csv(raw_final / "saturation_filter_row_counts.csv")
    bmp_counts = pd.read_csv(bmp_final / "saturation_filter_row_counts.csv")
    bias_summary = pd.read_csv(sat_dir / "standard_vs_intercept_corrected_summary.csv")
    raw_bias = bias_summary[bias_summary["run"] == "RAW saturation excluded"].copy()
    bmp_bias = bias_summary[bias_summary["run"] == "BMP saturation excluded"].copy()

    doc = Document()
    set_doc_style(doc)
    add_title(
        doc,
        "10 相机黑体炉标定技术报告",
        f"RAW / BMP 标定流程、容错处理、过曝剔除与最终拟合结果总结 · {date.today().isoformat()}",
    )

    add_callout(
        doc,
        "最终建议",
        "主标定建议采用 RAW 分支：自动 ROI + repeat 容错 + Cam01 暗场 + 坏点屏蔽 + 过曝条件剔除 + standard poly3。"
        "BMP 分支作为显示值标定分支可用，默认不减暗场；若只看指标，BMP 的残差偏置修正版略优。"
    )

    doc.add_heading("1. 标定数据与目标", level=1)
    doc.add_paragraph(
        "本次标定使用两天采集的黑体炉明场数据，温度覆盖 800-1500°C。"
        "每个工况理论上 10 个相机各拍 3 次，暗场只采集了 Camera01，作为十相机统一暗场参考。"
        "目标是得到每个相机、每个颜色通道从相机响应量 X 到等效黑体辐亮度 L_eff 的标定关系。"
    )
    add_table_from_rows(
        doc,
        ["项目", "内容"],
        [
            ["明场数据", "CalibrationCapture_20260522-23_800-1500_12bit16"],
            ["暗场数据", "CalibrationCapture_20260523171919_12bit16 / camera01_dark_images"],
            ["RAW 主结果", str(raw_final)],
            ["BMP 分支结果", str(bmp_final)],
            ["诊断汇总", str(sat_dir / "saturation_exclusion_summary.xlsx")],
        ],
        widths=[2200, 7160],
        font_size=8,
    )

    doc.add_heading("2. 总体处理流程", level=1)
    add_table_from_rows(
        doc,
        ["步骤", "处理内容", "输出/作用"],
        [
            ["1", "合并两天明场文件夹，删除 repeat4", "形成统一 800-1500°C 明场目录"],
            ["2", "按标定工况表过滤 planned 条件", "去掉非计划工况和多拍工况"],
            ["3", "repeat 容错", "某工况缺 1 次 repeat 时，用剩余 2 次平均"],
            ["4", "auto-anchor ROI", "高温图定位黑体炉中心，低温局部搜索或回退锚点"],
            ["5", "暗场与坏点处理", "按曝光匹配暗场；用暗场 median >=8192 构建坏点 mask"],
            ["6", "过曝诊断与剔除", "基于实际 ROI 检查 RAW/BMP 是否饱和，删除明确过曝条件"],
            ["7", "拟合与对比", "比较 standard / intercept-corrected，最终输出每相机每通道系数"],
        ],
        widths=[900, 3300, 5160],
        font_size=8,
    )

    doc.add_heading("3. 脚本实现概要", level=1)
    doc.add_paragraph("核心脚本和职责如下。")
    add_table_from_rows(
        doc,
        ["脚本", "实现内容"],
        [
            ["images_to_blackbody_measurements.py", "扫描 RAW/BMP 文件、解析温度/曝光/repeat、自动 ROI、repeat 容错、暗场匹配、坏点 mask、生成 measurement CSV。"],
            ["radiometric_calibration_rgb.py", "加载测量 CSV，积分光谱响应得到 L_eff，生成 fit table，执行 standard 与 intercept-corrected poly3 拟合并输出系数/图。"],
            ["export_monotonic_calibration.py", "后续 LUT/PCHIP 导出脚本，目前因过曝问题先搁置，等最终数据筛选确认后再继续。"],
            ["诊断脚本/诊断表", "对实际 auto ROI 做 RAW/BMP 饱和统计，输出按 repeat/channel 和 condition 的过曝诊断 CSV。"],
        ],
        widths=[3000, 6360],
        font_size=8,
    )

    doc.add_paragraph("最终主模型公式：")
    add_code_block(
        doc,
        "RAW standard:\n"
        "X_c = (DN_c - Dark_c) / exposure_s\n"
        "L_eff_c = a_c * X_c^3 + b_c * X_c^2 + c_c * X_c + d_c\n\n"
        "BMP direct-fit:\n"
        "X_c = DN_bmp,c / exposure_s\n"
        "L_eff_c = a_c * X_c^3 + b_c * X_c^2 + c_c * X_c + d_c",
    )

    doc.add_heading("4. 关键容错与质量控制", level=1)
    doc.add_heading("4.1 Repeat 容错", level=2)
    doc.add_paragraph(
        "每个工况期望 repeat1/2/3。若某个工况只缺一次 repeat，则不整体丢弃该工况，而是用可用 repeat 做平均。"
        "例如 800°C / 50 ms / Cam01 缺 repeat1 时，用 repeat2 和 repeat3 平均。"
    )

    doc.add_heading("4.2 自动 ROI", level=2)
    doc.add_paragraph(
        "由于换相机可能造成黑体炉位置偏移，固定 ROI 不可靠。脚本使用高温图建立稳定 anchor，"
        "再对每张图在 anchor 附近局部搜索热点；若低温图环境亮度干扰搜索，则回退到 anchor ROI。"
        "全部 ROI 坐标写入 roi_audit.csv，便于复查。"
    )

    doc.add_heading("4.3 暗场和坏点", level=2)
    doc.add_paragraph(
        "RAW 分支使用 Camera01 暗场作为十相机统一暗场。暗场按曝光匹配，允许 0.002 ms 的曝光舍入容差。"
        "坏点 mask 从同曝光暗场 RAW 的 median 图生成，阈值为 8192 DN；ROI 求均值时忽略这些像素。"
        "坏点处理对本次 120x120 ROI 的最终拟合影响很小，但作为长曝光暗场保护是有必要的。"
    )

    doc.add_heading("4.4 过曝剔除", level=2)
    doc.add_paragraph(
        "后续复查发现每个温度下高曝光尾部存在明显过曝。过曝会让响应曲线在高端被裁剪，"
        "把 poly3 拟合向错误方向拉偏，因此最终拟合必须先剔除过曝条件。"
    )
    add_table_from_rows(
        doc,
        ["分支", "平均原始条件", "平均保留", "平均剔除", "说明"],
        [
            ["RAW", raw_summary["avg_input_rows"].mean(), raw_summary["avg_kept_rows"].mean(), raw_summary["avg_excluded_rows"].mean(), "按实际 auto ROI 检查 RAW 饱和，并套用坏点 mask"],
            ["BMP", bmp_summary["avg_input_rows"].mean(), bmp_summary["avg_kept_rows"].mean(), bmp_summary["avg_excluded_rows"].mean(), "按实际 auto ROI 检查 8-bit BMP 是否达到 255 或接近 255"],
        ],
        widths=[1000, 1600, 1500, 1500, 3760],
        font_size=8,
    )
    add_picture(doc, asset_dir / "saturation_filter_counts.png", "图 1：RAW/BMP 每个相机保留与剔除的工况数量。")
    add_picture(doc, asset_dir / "saturation_exclusion_metric_comparison.png", "图 2：过曝剔除前后，RAW/BMP 的平均 RMSE 与 R2 对比。")

    doc.add_heading("5. 拟合方式选择", level=1)
    doc.add_paragraph(
        "前期比较了 linear、poly2、poly3、logpoly2，并对离群点删除做了试验。当前最稳的系数型模型仍然是 poly3。"
        "在过曝剔除后，R/G 通道的拟合质量显著提升，B 通道略有 RMSE 牺牲但仍保持很高 R2。"
    )
    add_dataframe_table(
        doc,
        raw_summary,
        headers=["通道", "相机数", "原 RMSE", "新 RMSE", "RMSE变化", "原R2", "新R2", "R2变化"],
        columns=["channel", "cameras", "old_rmse_mean", "new_rmse_mean", "delta_rmse_mean", "old_r2_mean", "new_r2_mean", "delta_r2_mean"],
        widths=[900, 900, 1300, 1300, 1300, 1200, 1200, 1260],
        font_size=8,
    )
    doc.paragraphs[-1].insert_paragraph_before("表 1：RAW 分支过曝剔除前后指标。")
    add_dataframe_table(
        doc,
        bmp_summary,
        headers=["通道", "相机数", "原 RMSE", "新 RMSE", "RMSE变化", "原R2", "新R2", "R2变化"],
        columns=["channel", "cameras", "old_rmse_mean", "new_rmse_mean", "delta_rmse_mean", "old_r2_mean", "new_r2_mean", "delta_r2_mean"],
        widths=[900, 900, 1300, 1300, 1300, 1200, 1200, 1260],
        font_size=8,
    )
    doc.paragraphs[-1].insert_paragraph_before("表 2：BMP 分支过曝剔除前后指标。")

    doc.add_heading("6. 残差偏置修正对比", level=1)
    doc.add_paragraph(
        "脚本同时输出 standard 与 intercept-corrected 两套结果。过曝剔除后的对比显示：RAW 分支 R/G 通道 standard 更好，"
        "B 通道偏置修正略好但改善很小；BMP 分支偏置修正在三通道上略有改善。"
    )
    add_dataframe_table(
        doc,
        raw_bias,
        headers=["通道", "standard RMSE", "修正 RMSE", "RMSE变化", "standard R2", "修正R2", "R2变化"],
        columns=["channel", "standard_rmse_mean", "intercept_rmse_mean", "delta_rmse_mean", "standard_r2_mean", "intercept_r2_mean", "delta_r2_mean"],
        widths=[900, 1500, 1500, 1300, 1400, 1400, 1360],
        font_size=8,
    )
    doc.paragraphs[-1].insert_paragraph_before("表 3：RAW 分支 standard vs residual-intercept corrected。")
    add_dataframe_table(
        doc,
        bmp_bias,
        headers=["通道", "standard RMSE", "修正 RMSE", "RMSE变化", "standard R2", "修正R2", "R2变化"],
        columns=["channel", "standard_rmse_mean", "intercept_rmse_mean", "delta_rmse_mean", "standard_r2_mean", "intercept_r2_mean", "delta_r2_mean"],
        widths=[900, 1500, 1500, 1300, 1400, 1400, 1360],
        font_size=8,
    )
    doc.paragraphs[-1].insert_paragraph_before("表 4：BMP 分支 standard vs residual-intercept corrected。")

    add_callout(
        doc,
        "模型选择结论",
        "RAW 最终建议用 standard poly3，不做残差偏置修正；BMP 如果追求指标可用 intercept-corrected，"
        "但为了与 RAW 形式统一，也可以保留 standard。当前主线仍以 RAW standard 作为最终推荐。"
    )

    doc.add_heading("7. 10 相机最终结果", level=1)
    doc.add_heading("7.1 RAW 主标定结果", level=2)
    doc.add_paragraph("RAW 主标定结果采用：过曝剔除 + 坏点处理 + standard poly3。下图为 10 相机三通道拟合总览。")
    add_picture(doc, asset_dir / "raw_saturation_excluded_fit_overview_part1.png", "图 3：RAW 过曝剔除后 Camera01-Camera05 拟合总览。")
    add_picture(doc, asset_dir / "raw_saturation_excluded_fit_overview_part2.png", "图 4：RAW 过曝剔除后 Camera06-Camera10 拟合总览。")

    doc.add_heading("7.2 BMP 直接拟合结果", level=2)
    doc.add_paragraph("BMP 分支使用 8-bit BMP ROI 均值直接拟合，默认不减暗场。该分支适合作为显示值/预览值的经验标定。")
    add_picture(doc, asset_dir / "bmp_saturation_excluded_fit_overview_part1.png", "图 5：BMP 过曝剔除后 Camera01-Camera05 拟合总览。")
    add_picture(doc, asset_dir / "bmp_saturation_excluded_fit_overview_part2.png", "图 6：BMP 过曝剔除后 Camera06-Camera10 拟合总览。")

    doc.add_heading("8. 文件输出位置", level=1)
    doc.add_paragraph(f"以下路径均相对于统一数据目录：{root}")
    add_table_from_rows(
        doc,
        ["类型", "路径"],
        [
            ["RAW 最终系数", rel(raw_final / "cameraXX" / "calibration_coefficients_standard.json", root)],
            ["BMP standard 系数", rel(bmp_final / "cameraXX" / "calibration_coefficients_standard.json", root)],
            ["BMP 修正系数", rel(bmp_final / "cameraXX" / "calibration_coefficients_intercept_corrected.json", root)],
            ["过曝诊断总表", rel(sat_dir / "saturation_exclusion_summary.xlsx", root)],
            ["脚本副本", str(root.parent / "calibration_scripts")],
        ],
        widths=[2200, 7160],
        font_size=8,
    )

    doc.add_page_break()
    doc.add_heading("附录 A：RAW 最终 standard poly3 系数", level=1)
    raw_coef = coefficient_rows(raw_final, corrected=False)
    for ch in ("R", "G", "B"):
        doc.add_heading(f"RAW {ch} 通道", level=2)
        df = raw_coef[raw_coef["channel"] == ch]
        add_dataframe_table(
            doc,
            df,
            headers=["相机", "通道", "a(x3)", "b(x2)", "c(x)", "d", "RMSE", "R2"],
            columns=["camera", "channel", "a_x3", "b_x2", "c_x", "d", "rmse", "r2"],
            widths=[1100, 800, 1300, 1300, 1300, 1300, 1100, 1160],
            font_size=7,
        )

    doc.add_page_break()
    doc.add_heading("附录 B：BMP 分支 intercept-corrected poly3 系数", level=1)
    bmp_coef = coefficient_rows(bmp_final, corrected=True)
    for ch in ("R", "G", "B"):
        doc.add_heading(f"BMP {ch} 通道", level=2)
        df = bmp_coef[bmp_coef["channel"] == ch]
        add_dataframe_table(
            doc,
            df,
            headers=["相机", "通道", "a(x3)", "b(x2)", "c(x)", "d", "RMSE", "R2"],
            columns=["camera", "channel", "a_x3", "b_x2", "c_x", "d", "rmse", "r2"],
            widths=[1100, 800, 1300, 1300, 1300, 1300, 1100, 1160],
            font_size=7,
        )

    doc.add_page_break()
    doc.add_heading("附录 C：后续建议", level=1)
    doc.add_paragraph(
        "1. 将过曝剔除规则固化进正式脚本参数，避免后续手工诊断。\n"
        "2. 采集每个相机自己的暗场，替换当前 Camera01 暗场复用方案。\n"
        "3. 在最终过曝筛选规则确认后，再继续 PCHIP/LUT 分支，并优先用所有非过曝工况节点导出部署用 LUT。\n"
        "4. 用独立验证图像检查标定后的温度反推误差，避免只在训练工况上评价。"
    )

    doc.save(output_docx)
    print(output_docx)


if __name__ == "__main__":
    main()
